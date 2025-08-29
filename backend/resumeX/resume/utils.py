# import os
# import io
# import PyPDF2
# from docx import Document

# def extract_text_from_file(file_obj):
#     filename = file_obj.name.lower()
#     file_obj.seek(0)
#     text = ""

#     if filename.endswith(".pdf"):
#         reader = PyPDF2.PdfReader(file_obj)
#         for page in reader.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     elif filename.endswith(".docx"):
#         doc = Document(file_obj)
#         text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
#     else:
#         text = ""

#     return text
import PyPDF2
from docx import Document

def extract_text_from_resume_file(file_obj):
    file_obj.seek(0)
    filename = file_obj.name.lower()
    text = ""

    if filename.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(file_obj)
        for page in pdf_reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + " "
    elif filename.endswith('.docx'):
        doc = Document(file_obj)
        text = " ".join([p.text for p in doc.paragraphs if p.text.strip()])
    else:
        # Optionally handle .doc files or others
        text = ""

    return text.lower()

